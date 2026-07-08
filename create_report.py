import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    """Sets the background color of a cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding (in dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_left_border(cell, color_hex="4F46E5", sz="36"):
    """Sets a left border on a cell and clears other borders to make a callout box."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    # Left border
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz) # 36 is 4.5 pt
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color_hex)
    tcBorders.append(left)
    
    # Clear top, bottom, right
    for border_name in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
        
    tcPr.append(tcBorders)

def add_callout(doc, text, title="NOTE"):
    """Creates a beautifully styled callout box in the Word document."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F3F4F6") # Very light grey
    set_left_border(cell, "4F46E5", "36") # Indigo left border
    set_cell_margins(cell, top=140, bottom=140, left=200, right=140)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    run_title = p.add_run(f"{title}: ")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(10.5)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(79, 70, 229)
    
    run_text = p.add_run(text)
    run_text.font.name = 'Arial'
    run_text.font.size = Pt(10)
    run_text.font.italic = True
    run_text.font.color.rgb = RGBColor(55, 65, 81)
    
    # Add an empty spacing paragraph after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(8)

def main():
    doc = Document()
    
    # Set document margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Define Palette Colors
    COLOR_INDIGO = RGBColor(79, 70, 229)   # Primary Accent (#4F46E5)
    COLOR_PURPLE = RGBColor(109, 40, 217)  # Secondary Accent (#6D28D9)
    COLOR_TEAL = RGBColor(13, 148, 136)    # Tertiary Accent (#0D9488)
    COLOR_CHARCOAL = RGBColor(55, 65, 81)  # Body Text (#374151)
    
    # Set Normal Style defaults
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_CHARCOAL
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(8)

    # ----------------------------------------------------
    # COVER PAGE / HEADER
    # ----------------------------------------------------
    
    # Insert Logo
    logo_path = os.path.join(os.getcwd(), 'static', 'images', 'logo_nexora_full.png')
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Pt(24)
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(3.2))
    else:
        print("Warning: Logo image not found at", logo_path)
        
    # Spacer
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(36)
    
    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(8)
    run_title = p_title.add_run("NEXORA.AI WORKSPACE MANUAL")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_INDIGO
    
    # Document Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(72)
    run_sub = p_sub.add_run("A Complete Project Report & User Guide for Advanced Document Intelligence, Conversational RAG, and Interactive Cognitive Analytics")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(107, 114, 128)
    
    # Document Metadata
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(12)
    run_meta1 = p_meta.add_run("Prepared for: ")
    run_meta1.font.bold = True
    run_meta1.font.size = Pt(10.5)
    run_meta2 = p_meta.add_run("Nexora.AI Workspace Users\n")
    run_meta2.font.size = Pt(10.5)
    
    run_meta3 = p_meta.add_run("Version: ")
    run_meta3.font.bold = True
    run_meta3.font.size = Pt(10.5)
    run_meta4 = p_meta.add_run("1.0 (Production Release)\n")
    run_meta4.font.size = Pt(10.5)
    
    run_meta5 = p_meta.add_run("Date: ")
    run_meta5.font.bold = True
    run_meta5.font.size = Pt(10.5)
    run_meta6 = p_meta.add_run("July 2026\n")
    run_meta6.font.size = Pt(10.5)
    
    # Page Break
    doc.add_page_break()

    # ----------------------------------------------------
    # SECTION 1: EXECUTIVE OVERVIEW
    # ----------------------------------------------------
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    r_h1 = h1.add_run("1. Executive Overview")
    r_h1.font.name = 'Arial'
    r_h1.font.size = Pt(18)
    r_h1.font.bold = True
    r_h1.font.color.rgb = COLOR_INDIGO
    
    p = doc.add_paragraph(
        "Nexora.AI is an advanced enterprise intelligence workspace that dynamically transforms unstructured data into semantic QA engines, interactive knowledge graphs, and automated action trackers. By indexing corporate archives—ranging from multi-page PDFs, Word documents, to raw text transcripts—Nexora provides teams with real-time cognitive analytics, dynamic clustering, and high-fidelity summaries."
    )
    
    p = doc.add_paragraph(
        "Unlike basic search portals or simple chatbot wrappers, Nexora implements a dual-discovery framework: a conversational chatbot utilizing Retrieval-Augmented Generation (RAG) working in unison with ten distinct visual analytics workspaces. This ensures that users can drill deep into document specific details while maintaining a global, structured overview of their document universe."
    )
    
    add_callout(doc, "All files uploaded to Nexora.AI are stored in a secure, isolated sandbox dedicated exclusively to your session. Your data is never used for LLM fine-tuning or external model training.", "SECURITY ARCHITECTURE")

    # Key highlights
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    r_h2 = h2.add_run("Key Value Propositions & Design Objectives")
    r_h2.font.name = 'Arial'
    r_h2.font.size = Pt(14)
    r_h2.font.bold = True
    r_h2.font.color.rgb = COLOR_PURPLE
    
    highlights = [
        ("Deep Context Comprehension", "Nexora reads beyond basic keyword strings, utilizing semantic sentence embeddings to map context, emotional sentiment, and core concepts across your files."),
        ("Zero Hallucination Guarantee", "The generative engine is strictly bounded to the uploaded document boundaries. If an answer cannot be verified within the corpus, the system explicitly alerts the user rather than fabricating data."),
        ("Automated Metadata Extraction", "Extracts complex named entities (entities like people, dates, organizations, and quantities) and reads document structure to construct dynamic tables and heatmaps."),
        ("Cross-Document Synthesis", "Enables inquiries spanning hundreds of pages across different file formats, synthesizing multi-file summaries and tracking conflicting statements."),
        ("Workspace Customization", "Features 6 premium visual themes (Lavender, Cyan, Sunset, Sherbet, Obsidian, Tropical) and full dark/light modes to support long analytical sessions."),
        ("Private Secure Sandbox", "Complete document isolation, isolated vector indexes, and session-specific databases guarantee total intellectual property safety.")
    ]
    
    for title, desc in highlights:
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.paragraph_format.space_after = Pt(4)
        run_bold = p_item.add_run(f"{title}: ")
        run_bold.bold = True
        run_bold.font.color.rgb = COLOR_TEAL
        p_item.add_run(desc)

    # Spacing
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ----------------------------------------------------
    # SECTION 2: STEP-BY-STEP USER GUIDE
    # ----------------------------------------------------
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    r_h1 = h1.add_run("2. Step-by-Step User Walkthrough")
    r_h1.font.name = 'Arial'
    r_h1.font.size = Pt(18)
    r_h1.font.bold = True
    r_h1.font.color.rgb = COLOR_INDIGO
    
    p = doc.add_paragraph(
        "To get started with Nexora.AI, follow this step-by-step workflow to upload your documents, explore information, and run advanced reports:"
    )
    
    steps = [
        ("Step 1: Upload Documents", "Locate the 'Upload Files' box in the left-hand sidebar. Drag and drop your PDFs or text documents directly into the zone, or click it to browse your computer. Once selected, click the 'Upload & Index' button. The platform will parse the text contents, partition them into semantic paragraphs, generate mathematical embeddings, and index them into a secure local vector database."),
        ("Step 2: Semantic Chat & QA", "Use the central conversation pane to ask questions in plain English. For example, you can type: 'Summarize the financial projections' or 'What are the main risk factors mentioned in these documents?'. Nexora searches the vector database, extracts the most relevant paragraphs, and feeds them as source parameters to the AI model to generate a precise response."),
        ("Step 3: Source Citation Inspection", "Every answer generated by Nexora includes citation badges linked to the source documents. Hover over or click a badge to display the Source Context Inspector. This sidebar opens the exact file page and highlights the matching sentence, letting you verify claims instantly without manual scroll-searching."),
        ("Step 4: Activating Explore Workspaces", "To unlock visual analytics, open the 'Explore Features' dropdown at the top navigation bar. Select any of the specialized analytical workspaces (such as Knowledge Graph, LexiScope, or Action Tracker). Click the 'Generate Report' button at the center of the viewport to process your document collection through that specific tool.")
    ]
    
    for title, desc in steps:
        p_step = doc.add_paragraph()
        p_step.paragraph_format.space_before = Pt(6)
        p_step.paragraph_format.space_after = Pt(4)
        run_bold = p_step.add_run(title)
        run_bold.bold = True
        run_bold.font.size = Pt(12)
        run_bold.font.color.rgb = COLOR_PURPLE
        
        p_desc = doc.add_paragraph(desc)
        p_desc.paragraph_format.left_indent = Inches(0.25)
        p_desc.paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ----------------------------------------------------
    # SECTION 3: THE TEN SUPERPOWERS (FEATURE DEEP DIVE)
    # ----------------------------------------------------
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    r_h1 = h1.add_run("3. The Ten Superpowers: Workspace Features")
    r_h1.font.name = 'Arial'
    r_h1.font.size = Pt(18)
    r_h1.font.bold = True
    r_h1.font.color.rgb = COLOR_INDIGO
    
    p = doc.add_paragraph(
        "Nexora's cognitive capabilities are distributed across ten specialized analytical tools. Each tool can be launched directly inside the embedded workspace overlay. By selecting specific files in your document list and clicking 'Generate Report' at the center, you activate these features:"
    )
    
    features = [
        ("Executive Report (One-Click Report Builder)", 
         "fa-file-invoice",
         "The Executive Report builder compiles a comprehensive executive summary of your documents. It uses AI to extract key milestones, background summaries, risk analyses, and operational recommendations. Users can select all or subset of chapters and export the compiled document straight as a formatted Word (.docx) or PDF file."),
         
        ("Action Item & Decision Tracker",
         "fa-circle-check",
         "The Action Tracker scans documents (especially project guidelines, meeting transcripts, and board notes) to automatically identify actionable tasks, deadlines, task owners, and major decision points. It lists them in a Kanban board layout where users can assign task statuses, view citations, and organize priorities."),
         
        ("TrendLens™ Live Verification",
         "fa-earth-americas",
         "TrendLens compares your document's internal assertions against live web intelligence. It runs automated search queries across global news feeds, reviews, tweets, and research articles to build a 'Doc vs. World' report. It categorizes claims as Confirmed (supported by web evidence), Contradicted (refuted by current reports), or Outdated."),
         
        ("Knowledge Graph",
         "fa-diagram-project",
         "The Knowledge Graph extracts concepts, entities (people, companies, locations), and topics from files, drawing them as a dynamic, interactive web. Powered by a force-directed layout, it reveals semantic links between documents, letting users spot hidden cross-document relationships immediately."),
         
        ("Cluster Universe",
         "fa-circle-nodes",
         "Cluster Universe plots all files in a 2D coordinate cloud. Using document embedding similarity, files with closely related contexts are grouped together. This provides a bird's-eye view of your document structure, sorting messy archives into logical semantic clusters without manual tagging."),
         
        ("TimelineWeave",
         "fa-timeline",
         "TimelineWeave extracts chronologically dated topics and arguments from your files and maps them against real-world external history. It fetches contemporary news reports and timelines, weaving your document's narrative alongside external historical events in a clean graphical timeline."),
         
        ("Visual Pulse Mood Board",
         "fa-images",
         "Visual Pulse reads the main subjects and thematic concepts of your files and searches the web for relevant high-fidelity images. It arranges these images into an interactive, visual mood-board gallery, providing visual context to supplement dense, text-based analytical files."),
         
        ("StatSonar (PulseGrid)",
         "fa-chart-pie",
         "StatSonar focuses on quantitative indexing and source validation. It reads numerical claims inside documents, lists them in card gauges, and builds a source-credibility heatmap. It groups web resources supporting your document's topics into credibility tiers (news, blogs, social chatter) to analyze reliability."),
         
        ("LexiScope Suite",
         "fa-flask-vial",
         "LexiScope is a comprehensive Natural Language Processing (NLP) dashboard. It computes overall sentence-level emotional sentiment, produces custom frequency word clouds, groups documents by keyword clusters, extracts named entities, and rates the text's linguistic readability score."),
         
        ("Interactive Mind Map",
         "fa-brain",
         "The Mind Map feature compiles your document's logical outline into an interactive tree of main ideas. Users can expand branches, collapse sub-themes, and click nodes to open detailed summaries. This radial concept tree is ideal for studying complex document structures in seconds.")
    ]
    
    for name, icon, desc in features:
        p_feat = doc.add_paragraph()
        p_feat.paragraph_format.space_before = Pt(8)
        p_feat.paragraph_format.space_after = Pt(2)
        run_name = p_feat.add_run(f"■ {name}")
        run_name.bold = True
        run_name.font.size = Pt(12)
        run_name.font.color.rgb = COLOR_PURPLE
        
        p_desc = doc.add_paragraph(desc)
        p_desc.paragraph_format.left_indent = Inches(0.2)
        p_desc.paragraph_format.space_after = Pt(10)

    doc.add_page_break()

    # ----------------------------------------------------
    # SECTION 4: TECHNICAL SYSTEM ARCHITECTURE
    # ----------------------------------------------------
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    r_h1 = h1.add_run("4. Technical System Architecture")
    r_h1.font.name = 'Arial'
    r_h1.font.size = Pt(18)
    r_h1.font.bold = True
    r_h1.font.color.rgb = COLOR_INDIGO
    
    p = doc.add_paragraph(
        "Nexora.AI uses a high-performance, modular software stack designed to support fast RAG search and dynamic web verification. The backend is written in Python using Flask and LangChain, while the front end utilizes a custom CSS glassmorphism UI with vanilla JavaScript."
    )
    
    # Let's add a small table to show the technology components
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    widths = [Inches(1.5), Inches(2.2), Inches(2.8)]
    headers = ["Layer", "Technologies Used", "Purpose / Functionality"]
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(hdr_cells[i], "4F46E5") # Indigo Header
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
        
    tech_data = [
        ("Application Server", "Flask, Python, api_router.py", "Handles API routing, user authentication, file uploads, and acts as the orchestrator between the front-end JS and analytical workers."),
        ("Orchestration & LLM", "LangChain, Groq API / Gemini", "Manages conversational memory, prompt templates, QA chain execution, and semantic extraction for analytical tasks."),
        ("Vector Store & Search", "ChromaDB, cosine similarity, embeddings", "Stores split paragraphs as mathematical vectors. Performs dense similarity search to fetch high-relevance chunks for context."),
        ("Web Intelligence", "httpx, web_augmentor.py, DuckDuckGo API", "Scans news feeds, web pages, and blogs asynchronously to retrieve live comparative references for TrendLens and StatSonar."),
        ("UI & Styling", "HTML5, CSS Custom Variables, glassmorphism", "Renders a modern, responsive dashboard with floating control bars, modal overlays, visual graphs, and dark/light color themes.")
    ]
    
    for row_idx, data in enumerate(tech_data, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            # Styling cell paragraphs
            p_cell = row_cells[col_idx].paragraphs[0]
            p_cell.paragraph_format.space_after = Pt(2)
            p_cell.paragraph_format.line_spacing = 1.1
            p_cell.runs[0].font.size = Pt(9.5)
            # Add thin borders and alternate row backgrounds
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=80, right=80)
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "F3F4F6") # Zebra striping
            else:
                set_cell_background(row_cells[col_idx], "FFFFFF")
                
    # Column width setting
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Spacing
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ----------------------------------------------------
    # SECTION 5: FREQUENTLY ASKED QUESTIONS (FAQ)
    # ----------------------------------------------------
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    r_h1 = h1.add_run("5. System FAQ & Troubleshooting")
    r_h1.font.name = 'Arial'
    r_h1.font.size = Pt(18)
    r_h1.font.bold = True
    r_h1.font.color.rgb = COLOR_INDIGO
    
    faqs = [
        ("Q: What is TrendLens and how does it work?",
         "A: TrendLens extracts core claims and topics from your uploaded documents and runs real-time web searches to check them against live internet resources. It compiles this evidence to show whether your documents are confirmed, contradicted, or outdated relative to current information."),
         
        ("Q: What is TimelineWeave?",
         "A: TimelineWeave automatically identifies the primary themes in your documents and graphs them chronologically alongside live web events, tweets, news, and milestones. It helps you see how your document's storyline aligns with real-world history."),
         
        ("Q: What is Visual Pulse?",
         "A: Visual Pulse generates an interactive image gallery and visual mood-board. It analyzes document topics and pairs them with relevant real-world images searched live from the web, providing visual context to text-based reports."),
         
        ("Q: What is StatSonar?",
         "A: StatSonar displays live metric counts, animated gauge cards, and source-credibility grids. It counts web references across news, blogs, and social feeds for your topics and groups them into credibility tiers to analyze source patterns."),
         
        ("Q: How does the AI search across my files?",
         "A: The system splits documents into paragraphs, embeds them into a semantic vector database, and matches your questions using dense cosine similarity to locate the most relevant sections before generating answers."),
         
        ("Q: Can I change colors or look?",
         "A: Yes! Click the gear icon in the top right of the dashboard to open 'Session Settings'. You can select from 6 themes (Lavender, Cyan, Sunset, Sherbet, Obsidian, Tropical) and toggle between dark and light appearance modes."),
         
        ("Q: How do I export my conversation?",
         "A: Use the floating PDF download button (represented by a file-download icon) in the top right of the chat area to download a beautifully formatted transcript of your session.")
    ]
    
    for q, a in faqs:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(6)
        p_q.paragraph_format.space_after = Pt(2)
        run_q = p_q.add_run(q)
        run_q.bold = True
        run_q.font.color.rgb = COLOR_PURPLE
        
        p_a = doc.add_paragraph(a)
        p_a.paragraph_format.left_indent = Inches(0.2)
        p_a.paragraph_format.space_after = Pt(10)
        
    # Final Footer Text
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(24)
    run_foot = p_foot.add_run("© 2026 Nexora.AI. All rights reserved. Confidential Document.")
    run_foot.font.size = Pt(9)
    run_foot.font.italic = True
    run_foot.font.color.rgb = RGBColor(156, 163, 175)

    # Save Document
    filename = "Nexora_AI_Project_Report.docx"
    doc.save(filename)
    print(f"Report created successfully at: {os.path.abspath(filename)}")

if __name__ == '__main__':
    main()

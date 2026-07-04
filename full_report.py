"""
full_report.py — Nexora AI  Full Report Generator
==================================================
Generates comprehensive PDF and DOCX reports from all (or selected)
indexed documents using the RAG pipeline to produce structured sections:
  • Executive Summary
  • Key Findings
  • Detailed Analysis (per-document)
  • Conclusion & Recommendations

Dependencies:
    pip install reportlab python-docx tavily-python pillow requests

Env vars:
    TAVILY_API_KEY      — required to enable the "Include Web Research" toggle
    NEXORA_LOGO_PATH    — optional explicit path to logo_nexora_full.png

Flask route wiring (in your app.py, inside /generate_report):
    use_web = bool(request.json.get("use_web", False))
    file_bytes, filename, mime, web_warning = generate_full_report(
        username, session_id, settings, selected_docs, fmt, use_web=use_web
    )
    resp = send_file(io.BytesIO(file_bytes), mimetype=mime,
                      download_name=filename, as_attachment=True)
    if web_warning:
        # Lets the frontend tell the user *why* "Include Web Research" didn't
        # add anything (e.g. missing TAVILY_API_KEY) instead of failing silently.
        resp.headers["X-Report-Web-Warning"] = web_warning
        resp.headers["Access-Control-Expose-Headers"] = "Content-Disposition, X-Report-Web-Warning"
    return resp
"""

from __future__ import annotations

import io
import os
import re
import sqlite3
import textwrap
from datetime import datetime
from typing import Optional

import chromadb
import requests
from io import BytesIO
from PIL import Image as PILImage

# ── Tavily (web research) ────────────────────────────────────────────────────
try:
    from tavily import TavilyClient
except Exception:
    TavilyClient = None  # feature degrades gracefully if package isn't installed

# ── ReportLab (PDF) ──────────────────────────────────────────────────────────
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

# ── python-docx (DOCX) ───────────────────────────────────────────────────────
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

# ── Project imports ───────────────────────────────────────────────────────────
from rag_logic import CHROMA_PATH, CHROMA_COLLECTION, get_llm

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────

BRAND_PRIMARY   = colors.HexColor("#6C63FF")   # Nexora violet
BRAND_SECONDARY = colors.HexColor("#A78BFA")
BRAND_DARK      = colors.HexColor("#1E1B2E")
BRAND_LIGHT     = colors.HexColor("#F5F3FF")
ACCENT_RULE     = colors.HexColor("#7C3AED")
WEB_ACCENT      = colors.HexColor("#22D3EE")   # cyan accent for web-sourced content

TAVILY_API_KEY  = os.getenv("TAVILY_API_KEY", "").strip()
WEB_MAX_RESULTS = 6      # how many web results Tavily returns per query
WEB_MAX_IMAGES  = 4      # how many images we embed in the report

# Candidate locations for the Nexora logo (first match wins)
LOGO_CANDIDATES = [
    os.getenv("NEXORA_LOGO_PATH", ""),
    "static/images/logo_nexora_full.png",
    "logo_nexora_full.png",
]

SECTIONS = [
    ("executive_summary",       "Executive Summary"),
    ("key_findings",            "Key Findings"),
    ("detailed_analysis",       "Detailed Analysis"),
    ("conclusion",              "Conclusion & Recommendations"),
]

DB_NAME = os.getenv("DB_NAME", "chat_history.db")


# ─────────────────────────────────────────────────────────────────────────────
# SHARED HELPERS — logo & image handling (used by PDF + DOCX builders)
# ─────────────────────────────────────────────────────────────────────────────

def _find_logo() -> Optional[str]:
    """Return the first existing logo path from LOGO_CANDIDATES, else None."""
    for path in LOGO_CANDIDATES:
        if path and os.path.isfile(path):
            return path
    return None


def _download_image_bytes(url: str, timeout: int = 6) -> Optional[bytes]:
    """Fetch an image from a URL. Returns None on any failure (never raises)."""
    try:
        resp = requests.get(url, timeout=timeout, headers={"User-Agent": "Mozilla/5.0"})
        resp.raise_for_status()
        # Validate it's actually a readable image before we trust it downstream
        PILImage.open(BytesIO(resp.content)).verify()
        return resp.content
    except Exception:
        return None


def _fit_image_reader(img_bytes: bytes, max_w: float, max_h: float):
    """Build a ReportLab Image flowable that fits within (max_w, max_h) while
    preserving aspect ratio. Returns None if the bytes aren't a valid image."""
    try:
        pil_img = PILImage.open(BytesIO(img_bytes))
        w, h = pil_img.size
        ratio = min(max_w / w, max_h / h)
        draw_w, draw_h = w * ratio, h * ratio
        return RLImage(BytesIO(img_bytes), width=draw_w, height=draw_h)
    except Exception:
        return None


# ─────────────────────────────────────────────────────────────────────────────
# STEP 1 — Retrieve all document text from ChromaDB
# ─────────────────────────────────────────────────────────────────────────────

def fetch_document_chunks(username: str, selected_docs: list[str]) -> dict[str, str]:
    """
    Pull every stored chunk for the given user (optionally filtered by filename).
    Returns { filename: concatenated_text }.
    """
    client = chromadb.PersistentClient(path=CHROMA_PATH)

    try:
        col = client.get_collection(CHROMA_COLLECTION)
    except Exception:
        return {}

    if selected_docs:
        if len(selected_docs) == 1:
            where: dict = {"$and": [{"username": username}, {"source": selected_docs[0]}]}
        else:
            where = {"$and": [{"username": username}, {"source": {"$in": selected_docs}}]}
    else:
        where = {"username": username}

    results = col.get(where=where, include=["documents", "metadatas"])
    doc_map: dict[str, list[tuple[int, str]]] = {}
    for doc_text, meta in zip(results.get("documents", []), results.get("metadatas", [])):
        source = meta.get("source", "Unknown")
        page   = int(meta.get("page", 0))
        doc_map.setdefault(source, []).append((page, doc_text))

    # Sort chunks by page order and join
    return {
        src: "\n\n".join(t for _, t in sorted(chunks))
        for src, chunks in doc_map.items()
    }


# ─────────────────────────────────────────────────────────────────────────────
# STEP 1B — Web research via Tavily (optional, ONE search call total)
# ─────────────────────────────────────────────────────────────────────────────

def _build_web_query(doc_map: dict[str, str]) -> str:
    """
    Derive a search query from the document set WITHOUT an extra LLM call —
    just the filenames (minus extensions) plus a short slice of the first
    document's text for topical grounding.
    """
    names = [re.sub(r"\.[a-zA-Z0-9]+$", "", n).replace("_", " ").replace("-", " ")
             for n in doc_map.keys()]
    query = " ".join(names[:4])[:120]

    if doc_map:
        first_text = next(iter(doc_map.values()))
        query += " " + first_text.strip()[:150]

    return query.strip()[:400] or "latest industry research and trends"


def _normalize_tavily_images(raw_images: list) -> list[dict]:
    """
    Tavily's `images` field is NOT stable across SDK/API versions:
      - Plain strings:              "https://…/img.jpg"
      - Dicts (image descriptions):  {"url": "https://…", "description": "…"}
    The previous implementation only accepted plain strings, which silently
    dropped 100% of results whenever Tavily returned the dict form — this
    was the root cause of "no images" in generated reports. Handle both.
    """
    normalized = []
    for img in raw_images or []:
        if isinstance(img, str) and img.strip():
            normalized.append({"url": img.strip(), "description": ""})
        elif isinstance(img, dict) and img.get("url"):
            normalized.append({
                "url":         img["url"].strip(),
                "description": (img.get("description") or "").strip(),
            })
        if len(normalized) >= WEB_MAX_IMAGES:
            break
    return normalized


def fetch_web_intel(doc_map: dict[str, str]) -> dict:
    """
    Single Tavily call that returns a dict that is ALWAYS populated (never
    None), so the caller — and the report itself — can tell the difference
    between "web research disabled" and "web research requested but failed",
    instead of the two looking identical downstream.

    Returned shape:
      {
        "enabled":  bool           — True only if we got usable results back
        "error":    str | None     — human-readable reason when enabled=False
        "query":    str
        "answer":   str            — Tavily's own synthesized answer
        "results":  [{title, url, content}, ...]
        "images":   [{url, description}, ...]
      }
    """
    base = {"enabled": False, "error": None, "query": "", "answer": "", "results": [], "images": []}

    if TavilyClient is None:
        base["error"] = "The 'tavily-python' package is not installed on the server."
        print(f"[Report] Web research skipped — {base['error']}")
        return base

    if not TAVILY_API_KEY:
        base["error"] = "TAVILY_API_KEY is not set in the server environment."
        print(f"[Report] Web research skipped — {base['error']}")
        return base

    query = _build_web_query(doc_map)
    base["query"] = query

    try:
        client = TavilyClient(api_key=TAVILY_API_KEY)
        resp = client.search(
            query=query,
            search_depth="advanced",
            include_answer=True,          # Tavily-generated synthesis — avoids a 2nd LLM call
            include_images=True,
            include_image_descriptions=True,   # ask explicitly for the {url, description} form
            max_results=WEB_MAX_RESULTS,
        )

        results = [
            {
                "title":   r.get("title", "Untitled Source"),
                "url":     r.get("url", ""),
                "content": (r.get("content") or "").strip(),
            }
            for r in resp.get("results", [])
            if r.get("url")
        ]
        images = _normalize_tavily_images(resp.get("images", []))

        if not results and not images and not resp.get("answer"):
            base["error"] = "Tavily returned an empty result set for this query."
            print(f"[Report] Web research call succeeded but returned nothing usable "
                  f"(query={query!r}).")
            return base

        base.update({
            "enabled": True,
            "answer":  (resp.get("answer") or "").strip(),
            "results": results,
            "images":  images,
        })
        print(f"[Report] Web research OK — {len(results)} sources, {len(images)} images.")
        return base

    except Exception as e:
        base["error"] = f"Tavily search failed: {e}"
        print(f"[Report] {base['error']} — continuing doc-only.")
        return base


def _web_intel_to_context(web_intel: dict) -> str:
    """Flatten web_intel into a compact text block for the LLM prompt."""
    if not web_intel or not web_intel.get("enabled"):
        return ""
    parts = []
    if web_intel.get("answer"):
        parts.append(f"Web Synthesis: {web_intel['answer']}")
    for r in web_intel.get("results", [])[:WEB_MAX_RESULTS]:
        snippet = r["content"][:400]
        parts.append(f"[{r['title']}]({r['url']}): {snippet}")
    return "\n\n".join(parts)


# ─────────────────────────────────────────────────────────────────────────────
# STEP 2 — LLM section generation
# ─────────────────────────────────────────────────────────────────────────────
def _call_llm(prompt: str, session_id: str, settings: dict) -> str:
    """Uses dedicated GROQ_API_KEY_FULL_REPORT, falls back to router."""
    import os
    from api_router import call_llm_with_fallback
    dedicated = os.getenv("GROQ_API_KEY_FULL_REPORT", "").strip()
    if dedicated:
        try:
            from langchain_groq import ChatGroq
            llm = ChatGroq(
                groq_api_key=dedicated,
                model_name=settings.get("model_name", "llama-3.3-70b-versatile"),
                temperature=settings.get("temperature", 0.0),
                max_tokens=2048,
            )
            return llm.invoke(prompt).content.strip()
        except Exception as e:
            print(f"[Report] Dedicated key failed ({e}), using router.")
    return call_llm_with_fallback(prompt, {**settings, "max_tokens": 2048})


def generate_section(
    section_key: str,
    section_title: str,
    doc_map: dict[str, str],
    session_id: str,
    settings: dict,
    web_context: str = "",
) -> str:
    from token_utils import trim_to_budget
    combined_context = ""
    for fname, text in doc_map.items():
        # ← KEY CHANGE: 1 200 tokens per doc instead of 4 000 chars
        snippet = trim_to_budget(text, 1_200)
        combined_context += f"\\n\\n--- Document: {fname} ---\\n{snippet}"

    # Trim the whole doc context to 6 000 tokens max
    combined_context = trim_to_budget(combined_context, 6_000)

    # Web research (if enabled) gets its own smaller budget so total prompt
    # size — and therefore LLM cost — stays predictable. No extra LLM calls
    # are made for the web content; it's folded into this same section call.
    web_block = ""
    if web_context:
        trimmed_web = trim_to_budget(web_context, 1_500)
        web_block = f"\\n\\nADDITIONAL WEB RESEARCH (use to add current, external context; " \
                    f"clearly complements but never contradicts the documents):\\n{trimmed_web}"

    # Shared framing so all 4 calls read as one consistent, senior-analyst voice
    # rather than 4 independently-styled LLM outputs stitched together.
    persona = (
        "You are a senior research analyst producing one section of an enterprise "
        "intelligence briefing for executive readers. Write in a precise, confident, "
        "analytical tone — no filler, no generic statements, no first-person, no "
        "restating the prompt. Ground every claim in the DOCUMENTS provided; never "
        "invent figures. Use plain text only (no markdown headers like ### or tables)."
    )
    web_rule = (
        " If ADDITIONAL WEB RESEARCH is provided, you may use it to add current, "
        "external context, but it must never contradict the documents, and every "
        "external figure or claim must be prefixed with 'External:' so readers can "
        "tell document-sourced content from web-sourced content."
    ) if web_block else ""

    prompts = {
        "executive_summary": f"""{persona}{web_rule}

Write the Executive Summary section (3-5 tight paragraphs, no bullets). Open with the
single most important takeaway. Then cover market size/scale, the key behavioral or
structural shift observed, and the strategic implication for a decision-maker reading
this briefing. Every paragraph must contain at least one concrete number from the
documents.

DOCUMENTS:{combined_context}{web_block}

Executive Summary only, plain text, no heading line:""",

        "key_findings": f"""{persona}{web_rule}

Write the Key Findings section: 8-12 findings, each as its own line in the exact
format "N. **Short Title**: one-sentence finding with a specific figure or fact."
Order findings from most strategically significant to least. Do not restate the same
figure twice. {'Include no more than 2 findings sourced from external web research, each still following the same "N. **Title**: fact" format with the fact itself prefixed "External:".' if web_block else ''}

DOCUMENTS:{combined_context}{web_block}

Key Findings only, plain text, no heading line:""",

        "detailed_analysis": f"""{persona}{web_rule}

Write the Detailed Analysis section using exactly these five labeled sub-sections,
each as its own paragraph starting with the label followed by a colon, in this order:
"Content Overview:", "Main Themes:", "Data & Metrics:", "Risks:", "Strengths:"{' and finally "External Market Context:"' if web_block else ''}.
Each sub-section should be 2-4 sentences of substantive analysis, not a list of topics.

DOCUMENTS:{combined_context}{web_block}

Detailed Analysis only, plain text, no heading line:""",

        "conclusion": f"""{persona}{web_rule}

Write two parts. First, a Conclusion of 2-3 paragraphs synthesizing the overall
narrative and what it means going forward. Second, on a new paragraph starting with
the exact line "Recommendations:", give 5-8 recommendations, each its own line in the
format "N. **Action-oriented title**: one sentence on why it matters and what to do."
Order by priority.{' Factor in the external web research where it strengthens a recommendation.' if web_block else ''}

DOCUMENTS:{combined_context}{web_block}

Conclusion and Recommendations only, plain text, no heading line:""",
    }

    prompt = prompts.get(section_key, f"{persona}\\n\\nSummarise these documents:\\n{combined_context}{web_block}")
    return _call_llm(prompt, session_id, settings)


def generate_all_sections(
    doc_map: dict[str, str],
    session_id: str,
    settings: dict,
    web_intel: Optional[dict] = None,
) -> dict[str, str]:
    """Generate content for every report section. If web_intel is provided,
    its content is merged into each of the same 4 LLM calls — this NEVER
    adds extra LLM hits, it just enriches the existing per-section prompts."""
    web_context = _web_intel_to_context(web_intel) if web_intel else ""
    result = {}
    for key, title in SECTIONS:
        result[key] = generate_section(key, title, doc_map, session_id, settings, web_context)
    return result


# ─────────────────────────────────────────────────────────────────────────────
# STEP 3 — PDF generation (ReportLab)
# ─────────────────────────────────────────────────────────────────────────────

def _pdf_styles():
    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        name="NexoraTitle",
        fontSize=26, leading=32, spaceAfter=6,
        textColor=BRAND_PRIMARY, fontName="Helvetica-Bold",
        alignment=1,
    ))
    styles.add(ParagraphStyle(
        name="NexoraSubtitle",
        fontSize=13, leading=18, spaceAfter=4,
        textColor=BRAND_SECONDARY, fontName="Helvetica",
        alignment=1,
    ))
    styles.add(ParagraphStyle(
        name="NexoraMeta",
        fontSize=10, leading=14, spaceAfter=2,
        textColor=colors.HexColor("#888888"), fontName="Helvetica",
        alignment=1,
    ))
    styles.add(ParagraphStyle(
        name="NexoraH1",
        fontSize=16, leading=22, spaceBefore=18, spaceAfter=8,
        textColor=BRAND_PRIMARY, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        name="NexoraH2",
        fontSize=13, leading=18, spaceBefore=12, spaceAfter=6,
        textColor=BRAND_SECONDARY, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        name="NexoraBody",
        fontSize=10, leading=15, spaceAfter=8,
        textColor=BRAND_DARK, fontName="Helvetica",
    ))
    styles.add(ParagraphStyle(
        name="NexoraCallout",
        fontSize=10, leading=15, spaceAfter=8,
        textColor=colors.HexColor("#4C1D95"),
        backColor=BRAND_LIGHT,
        fontName="Helvetica",
        leftIndent=12, rightIndent=12,
        spaceBefore=6,
        borderPad=8,
    ))
    styles.add(ParagraphStyle(
        name="NexoraBullet",
        fontSize=10, leading=15, spaceAfter=4,
        textColor=BRAND_DARK, fontName="Helvetica",
        leftIndent=20, bulletIndent=8,
    ))
    styles.add(ParagraphStyle(
        name="NexoraWebLink",
        fontSize=9.5, leading=14, spaceAfter=6,
        textColor=colors.HexColor("#0E7490"), fontName="Helvetica",
        leftIndent=4,
    ))
    styles.add(ParagraphStyle(
        name="NexoraCaption",
        fontSize=8.5, leading=11, spaceAfter=2,
        textColor=colors.HexColor("#888888"), fontName="Helvetica-Oblique",
        alignment=1,
    ))
    styles.add(ParagraphStyle(
        name="NexoraKicker",
        fontSize=9.5, leading=13, spaceAfter=8,
        textColor=WEB_ACCENT, fontName="Helvetica-Bold",
        alignment=1,
    ))
    styles.add(ParagraphStyle(
        name="NexoraStatValue",
        fontSize=15, leading=18,
        textColor=BRAND_PRIMARY, fontName="Helvetica-Bold",
        alignment=1,
    ))
    styles.add(ParagraphStyle(
        name="NexoraStatLabel",
        fontSize=8, leading=10, spaceAfter=0,
        textColor=colors.HexColor("#777777"), fontName="Helvetica",
        alignment=1,
    ))
    styles.add(ParagraphStyle(
        name="NexoraSectionBadge",
        fontSize=9, leading=11,
        textColor=colors.white, fontName="Helvetica-Bold",
        alignment=1,
    ))
    styles.add(ParagraphStyle(
        name="NexoraH1Num",
        fontSize=16, leading=22, spaceBefore=18, spaceAfter=8,
        textColor=BRAND_PRIMARY, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        name="NexoraWarning",
        fontSize=9, leading=13, spaceAfter=6,
        textColor=colors.HexColor("#92400E"),
        backColor=colors.HexColor("#FEF3C7"),
        fontName="Helvetica-Oblique",
        leftIndent=10, rightIndent=10, spaceBefore=4, borderPad=6,
    ))
    return styles


def _add_page_number(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#999999"))
    page_num = canvas.getPageNumber()
    canvas.drawCentredString(letter[0] / 2, 0.5 * inch, f"Nexora.AI — Confidential Report  |  Page {page_num}")
    canvas.restoreState()


def _clean_text_for_pdf(text: str) -> str:
    """Strip markdown syntax so ReportLab doesn't choke."""
    text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.*?)\*", r"<i>\1</i>", text)
    text = re.sub(r"#+\s+", "", text)          # headings
    text = re.sub(r"_{2,}", "", text)           # horizontal rules
    text = text.replace("&", "&amp;").replace("<b>", "<b>").replace("</b>", "</b>")
    return text


def _section_to_pdf_flowables(title: str, body: str, styles, index: Optional[int] = None,
                               highlight: bool = False) -> list:
    """Convert a section title + body text into ReportLab flowables.

    index      — if given, renders a small "01" style numbered badge before
                 the title, matching a newsletter/analytical-briefing layout.
    highlight  — if True (used for Executive Summary), wraps each paragraph
                 in the tinted NexoraCallout panel instead of plain body text.
    """
    flowables = []
    flowables.append(HRFlowable(width="100%", thickness=2, color=ACCENT_RULE, spaceAfter=6))

    if index is not None:
        badge_tbl = Table(
            [[Paragraph(f"{index:02d}", styles["NexoraSectionBadge"]),
              Paragraph(title, styles["NexoraH1Num"])]],
            colWidths=[0.36 * inch, 5.6 * inch],
        )
        badge_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, 0), BRAND_PRIMARY),
            ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN",      (0, 0), (0, 0), "CENTER"),
            ("TOPPADDING",    (0, 0), (0, 0), 5),
            ("BOTTOMPADDING", (0, 0), (0, 0), 5),
            ("LEFTPADDING",   (1, 0), (1, 0), 10),
        ]))
        flowables.append(badge_tbl)
    else:
        flowables.append(Paragraph(title, styles["NexoraH1"]))

    flowables.append(Spacer(1, 6))

    para_style = styles["NexoraCallout"] if highlight else styles["NexoraBody"]

    # Sub-heading labels the "Detailed Analysis" prompt is instructed to emit,
    # e.g. "Data & Metrics: ..." — render the label itself as a mini heading.
    label_re = re.compile(r"^(Content Overview|Main Themes|Data & Metrics|Risks|Strengths|"
                           r"External Market Context)\s*:\s*(.*)$", re.IGNORECASE | re.DOTALL)

    for para in body.split("\n\n"):
        para = para.strip()
        if not para:
            continue

        if para.lower().startswith("recommendations:"):
            flowables.append(Spacer(1, 6))
            flowables.append(Paragraph("Recommendations", styles["NexoraH2"]))
            para = para.split(":", 1)[1].strip()
            if not para:
                continue

        m = label_re.match(para)
        if m:
            flowables.append(Paragraph(m.group(1), styles["NexoraH2"]))
            cleaned = _clean_text_for_pdf(m.group(2).strip())
            if cleaned:
                flowables.append(Paragraph(cleaned, styles["NexoraBody"]))
            continue

        # Numbered list items ("N. **Title**: detail")
        if re.match(r"^\d+\.", para):
            lines = para.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                cleaned = _clean_text_for_pdf(re.sub(r"^\d+\.\s*", "", line))
                flowables.append(Paragraph(f"• {cleaned}", styles["NexoraBullet"]))
            flowables.append(Spacer(1, 4))
            continue

        # Sub-headings (markdown-style, in case the model still emits them)
        if para.startswith("#"):
            cleaned = re.sub(r"#+\s+", "", para)
            flowables.append(Paragraph(cleaned, styles["NexoraH2"]))
            continue

        cleaned = _clean_text_for_pdf(para)
        flowables.append(Paragraph(cleaned, para_style))
        if highlight:
            flowables.append(Spacer(1, 4))

    flowables.append(Spacer(1, 10))
    return flowables


def build_pdf_report(
    username: str,
    doc_map: dict[str, str],
    sections: dict[str, str],
    selected_docs: list[str],
    web_intel: Optional[dict] = None,
) -> bytes:
    """Assemble the full PDF and return raw bytes."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.85 * inch,
        leftMargin=0.85 * inch,
        topMargin=1 * inch,
        bottomMargin=0.85 * inch,
    )
    styles = _pdf_styles()
    story  = []
    web_on = bool(web_intel and web_intel.get("enabled"))

    # ── Cover Page ────────────────────────────────────────────────────────────
    logo_path = _find_logo()
    have_logo = False
    if logo_path:
        try:
            with open(logo_path, "rb") as f:
                logo_flowable = _fit_image_reader(f.read(), max_w=1.9 * inch, max_h=0.85 * inch)
            if logo_flowable:
                logo_flowable.hAlign = "CENTER"
                story.append(Spacer(1, 0.55 * inch))
                story.append(logo_flowable)
                story.append(Spacer(1, 0.25 * inch))
                have_logo = True
        except Exception:
            pass
    if not have_logo:
        story.append(Spacer(1, 1.1 * inch))

    story.append(Paragraph("ENTERPRISE INTELLIGENCE BRIEFING", styles["NexoraKicker"]))
    # Only render the big text wordmark as a fallback when there's no logo
    # image — the logo asset already carries the "Nexora.AI" wordmark, so
    # showing both was producing the duplicated brand name/logo on the cover.
    if not have_logo:
        story.append(Paragraph("NEXORA.AI", styles["NexoraTitle"]))
        story.append(Spacer(1, 0.1 * inch))
    story.append(HRFlowable(width="60%", thickness=2, color=BRAND_SECONDARY, hAlign="CENTER"))
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph("Comprehensive Document Intelligence Report", styles["NexoraSubtitle"]))
    story.append(Spacer(1, 0.35 * inch))

    # KPI stat strip — quick-glance metadata, newsletter/analytics-briefing style
    doc_names  = list(doc_map.keys())
    web_status = "Enabled" if web_on else ("Attempted" if (web_intel and web_intel.get("error")) else "Off")
    stat_cells = [
        Paragraph(str(len(doc_names)), styles["NexoraStatValue"]),
        Paragraph(datetime.now().strftime("%b %d, %Y"), styles["NexoraStatValue"]),
        Paragraph(web_status, styles["NexoraStatValue"]),
    ]
    stat_labels = [
        Paragraph("SOURCE DOCUMENTS", styles["NexoraStatLabel"]),
        Paragraph("GENERATED", styles["NexoraStatLabel"]),
        Paragraph("WEB RESEARCH", styles["NexoraStatLabel"]),
    ]
    stat_tbl = Table([stat_cells, stat_labels], colWidths=[2.0 * inch] * 3)
    stat_tbl.setStyle(TableStyle([
        ("ALIGN",  (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LINEBEFORE", (1, 0), (1, -1), 0.75, colors.HexColor("#DDDDDD")),
        ("LINEBEFORE", (2, 0), (2, -1), 0.75, colors.HexColor("#DDDDDD")),
        ("TOPPADDING",    (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(stat_tbl)
    story.append(Spacer(1, 0.25 * inch))

    story.append(Paragraph(f"Prepared for: {username}", styles["NexoraMeta"]))
    if web_intel and web_intel.get("error"):
        story.append(Spacer(1, 6))
        story.append(Paragraph(
            f"⚠ Web research was requested but unavailable — {web_intel['error']} "
            f"This report was generated using source documents only.",
            styles["NexoraWarning"],
        ))
    story.append(Spacer(1, 0.2 * inch))

    # Document list table
    table_data = [["#", "Document Name"]]
    for i, name in enumerate(doc_names, 1):
        table_data.append([str(i), name])

    tbl = Table(table_data, colWidths=[0.4 * inch, 5.6 * inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0),  BRAND_PRIMARY),
        ("TEXTCOLOR",     (0, 0), (-1, 0),  colors.white),
        ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS",(0, 1), (-1, -1), [colors.white, BRAND_LIGHT]),
        ("GRID",          (0, 0), (-1, -1), 0.5, colors.HexColor("#DDDDDD")),
        ("ALIGN",         (0, 0), (0, -1),  "CENTER"),
        ("TOPPADDING",    (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(tbl)
    story.append(PageBreak())

    # ── Report Sections ───────────────────────────────────────────────────────
    for idx, (key, title) in enumerate(SECTIONS, 1):
        body = sections.get(key, "Section not generated.")
        highlight = (key == "executive_summary")
        story.extend(_section_to_pdf_flowables(title, body, styles, index=idx, highlight=highlight))
        story.append(PageBreak())

    # ── Web Research Highlights (newsletter-style image gallery) ───────────────
    if web_on and web_intel.get("images"):
        gallery_rows = []
        row_imgs, row_caps = [], []
        for img in web_intel["images"][:WEB_MAX_IMAGES]:
            img_bytes = _download_image_bytes(img["url"])
            if not img_bytes:
                continue
            flow_img = _fit_image_reader(img_bytes, max_w=2.6 * inch, max_h=1.7 * inch)
            if not flow_img:
                continue
            caption = img.get("description") or "Web research"
            row_imgs.append(flow_img)
            row_caps.append(Paragraph(textwrap.shorten(caption, width=70, placeholder="…"),
                                       styles["NexoraCaption"]))
            if len(row_imgs) == 2:
                gallery_rows.append(row_imgs)
                gallery_rows.append(row_caps)
                row_imgs, row_caps = [], []
        if row_imgs:
            row_imgs.append("")
            row_caps.append("")
            gallery_rows.append(row_imgs)
            gallery_rows.append(row_caps)

        if gallery_rows:
            story.append(HRFlowable(width="100%", thickness=2, color=WEB_ACCENT, spaceAfter=6))
            story.append(Paragraph("Web Research Highlights", styles["NexoraH1"]))
            story.append(Spacer(1, 6))
            gtbl = Table(gallery_rows, colWidths=[2.8 * inch, 2.8 * inch])
            gtbl.setStyle(TableStyle([
                ("ALIGN",   (0, 0), (-1, -1), "CENTER"),
                ("VALIGN",  (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING",    (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(gtbl)
            story.append(Spacer(1, 10))
            story.append(PageBreak())

    # ── Web Sources ──────────────────────────────────────────────────────────
    if web_on and web_intel.get("results"):
        story.append(HRFlowable(width="100%", thickness=2, color=WEB_ACCENT, spaceAfter=6))
        story.append(Paragraph("Web Sources & Further Reading", styles["NexoraH1"]))
        for i, r in enumerate(web_intel["results"], 1):
            story.append(Paragraph(f"{i}. <b>{r['title']}</b> — {r['url']}", styles["NexoraWebLink"]))
        story.append(PageBreak())

    # ── Sources Footer ────────────────────────────────────────────────────────
    story.append(HRFlowable(width="100%", thickness=2, color=ACCENT_RULE, spaceAfter=6))
    story.append(Paragraph("Source Documents", styles["NexoraH1"]))
    for i, name in enumerate(doc_names, 1):
        story.append(Paragraph(f"{i}. {name}", styles["NexoraBody"]))

    doc.build(story, onFirstPage=_add_page_number, onLaterPages=_add_page_number)
    return buffer.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# STEP 4 — DOCX generation (python-docx)
# ─────────────────────────────────────────────────────────────────────────────

def _set_docx_heading_color(paragraph, hex_color: str):
    """Apply a custom colour to a heading paragraph's runs."""
    rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(*rgb)


def _shade_docx_paragraph(paragraph, hex_color: str):
    """Apply a light background shade to a paragraph — used to mimic the
    tinted callout-box treatment ReportLab gives the Executive Summary."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _add_horizontal_rule(doc: DocxDocument):
    """Insert a thin coloured paragraph border as a visual divider."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "7C3AED")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)


_DOCX_LABEL_RE = re.compile(
    r"^(Content Overview|Main Themes|Data & Metrics|Risks|Strengths|External Market Context)"
    r"\s*:\s*(.*)$", re.IGNORECASE | re.DOTALL,
)


def _section_to_docx(doc: DocxDocument, title: str, body: str, index: Optional[int] = None,
                      highlight: bool = False):
    """Write one section into the DOCX.

    index      — numbers the section heading ("01  Executive Summary") to
                 match the newsletter/analytical-briefing layout used in the PDF.
    highlight  — shades each paragraph's background for Executive Summary,
                 mirroring the PDF's tinted callout treatment.
    """
    _add_horizontal_rule(doc)
    heading_text = f"{index:02d}   {title}" if index is not None else title
    h = doc.add_heading(heading_text, level=1)
    _set_docx_heading_color(h, "6C63FF")
    h.paragraph_format.space_before = Pt(14)

    for para_text in body.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue

        if para_text.lower().startswith("recommendations:"):
            doc.add_heading("Recommendations", level=2)
            para_text = para_text.split(":", 1)[1].strip()
            if not para_text:
                continue

        label_match = _DOCX_LABEL_RE.match(para_text)
        if label_match:
            sh = doc.add_heading(label_match.group(1), level=2)
            _set_docx_heading_color(sh, "A78BFA")
            rest = label_match.group(2).strip()
            if rest:
                p = doc.add_paragraph(rest)
                p.runs[0].font.size = Pt(10.5)
                if highlight:
                    _shade_docx_paragraph(p, "F5F3FF")
            continue

        # Sub-headings
        if para_text.startswith("#"):
            clean = re.sub(r"#+\s+", "", para_text)
            sh = doc.add_heading(clean, level=2)
            _set_docx_heading_color(sh, "A78BFA")
            continue

        # Numbered list lines
        if re.match(r"^\d+\.", para_text):
            for line in para_text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                p = doc.add_paragraph(style="List Bullet")
                # Strip leading "N." if present
                clean_line = re.sub(r"^\d+\.\s*", "", line)
                # Handle bold markdown
                parts = re.split(r"\*\*(.*?)\*\*", clean_line)
                for j, part in enumerate(parts):
                    run = p.add_run(part)
                    if j % 2 == 1:
                        run.bold = True
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        parts = re.split(r"\*\*(.*?)\*\*", para_text)
        for j, part in enumerate(parts):
            run = p.add_run(part)
            run.font.size = Pt(10.5)
            if j % 2 == 1:
                run.bold = True
        if highlight:
            _shade_docx_paragraph(p, "F5F3FF")
            p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()  # breathing space


def build_docx_report(
    username: str,
    doc_map: dict[str, str],
    sections: dict[str, str],
    selected_docs: list[str],
    web_intel: Optional[dict] = None,
) -> bytes:
    """Assemble the full DOCX and return raw bytes."""
    doc = DocxDocument()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin  = section.bottom_margin = Inches(1)

    web_on = bool(web_intel and web_intel.get("enabled"))

    # ── Cover ─────────────────────────────────────────────────────────────────
    logo_path = _find_logo()
    have_logo = False
    if logo_path:
        try:
            logo_p = doc.add_paragraph()
            logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_p.add_run().add_picture(logo_path, width=Inches(1.7))
            have_logo = True
        except Exception:
            pass

    kicker = doc.add_paragraph("ENTERPRISE INTELLIGENCE BRIEFING")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.runs[0].font.size = Pt(10)
    kicker.runs[0].font.bold = True
    kicker.runs[0].font.color.rgb = RGBColor(0x22, 0xD3, 0xEE)

    # Only show the big text wordmark when there's no logo image — the logo
    # asset already contains the "Nexora.AI" wordmark, so showing both was
    # producing the duplicated brand name/logo seen on the previous cover.
    if not have_logo:
        cover_title = doc.add_heading("NEXORA.AI", 0)
        cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_docx_heading_color(cover_title, "6C63FF")

    sub = doc.add_paragraph("Comprehensive Document Intelligence Report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0xA7, 0x8B, 0xFA)

    doc.add_paragraph()

    # KPI metadata row — quick-glance stats, newsletter/analytics-briefing style
    doc_names_preview = list(doc_map.keys())
    web_status = "Enabled" if web_on else ("Attempted" if (web_intel and web_intel.get("error")) else "Off")
    stat_tbl = doc.add_table(rows=2, cols=3)
    stat_tbl.autofit = True
    stat_values = [str(len(doc_names_preview)), datetime.now().strftime("%b %d, %Y"), web_status]
    stat_labels = ["SOURCE DOCUMENTS", "GENERATED", "WEB RESEARCH"]
    for col, (val, lbl) in enumerate(zip(stat_values, stat_labels)):
        vcell = stat_tbl.rows[0].cells[col]
        vcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vrun = vcell.paragraphs[0].add_run(val)
        vrun.font.size = Pt(15)
        vrun.font.bold = True
        vrun.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)
        lcell = stat_tbl.rows[1].cells[col]
        lcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        lrun = lcell.paragraphs[0].add_run(lbl)
        lrun.font.size = Pt(7.5)
        lrun.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()

    meta1 = doc.add_paragraph(f"Prepared for: {username}")
    meta1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta1.runs[0].font.size = Pt(10)
    meta1.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    if web_intel and web_intel.get("error"):
        warn = doc.add_paragraph(
            f"⚠ Web research was requested but unavailable — {web_intel['error']} "
            f"This report was generated using source documents only."
        )
        warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        warn.runs[0].font.size = Pt(9)
        warn.runs[0].italic = True
        warn.runs[0].font.color.rgb = RGBColor(0x92, 0x40, 0x0E)

    doc.add_paragraph()
    _add_horizontal_rule(doc)

    # Document index table
    doc_names = list(doc_map.keys())
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Document Name"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "6C63FF")
        shd.set(qn("w:val"), "clear")
        cell._tc.tcPr.append(shd)

    for i, name in enumerate(doc_names, 1):
        row = tbl.add_row().cells
        row[0].text = str(i)
        row[1].text = name

    doc.add_page_break()

    # ── Sections ─────────────────────────────────────────────────────────────
    for idx, (key, title) in enumerate(SECTIONS, 1):
        body = sections.get(key, "Section not generated.")
        highlight = (key == "executive_summary")
        _section_to_docx(doc, title, body, index=idx, highlight=highlight)
        doc.add_page_break()

    # ── Web Research Highlights (newsletter-style image gallery) ───────────────
    if web_on and web_intel.get("images"):
        downloaded = []  # [(img_bytes, caption)]
        for img in web_intel["images"][:WEB_MAX_IMAGES]:
            img_bytes = _download_image_bytes(img["url"])
            if img_bytes:
                downloaded.append((img_bytes, img.get("description") or "Web research"))

        if downloaded:
            _add_horizontal_rule(doc)
            gh = doc.add_heading("Web Research Highlights", level=1)
            _set_docx_heading_color(gh, "22D3EE")

            gtbl = doc.add_table(rows=0, cols=2)
            for i in range(0, len(downloaded), 2):
                img_row = gtbl.add_row().cells
                cap_row = gtbl.add_row().cells
                for j, (img_bytes, caption) in enumerate(downloaded[i:i + 2]):
                    try:
                        run = img_row[j].paragraphs[0].add_run()
                        run.add_picture(BytesIO(img_bytes), width=Inches(2.9))
                        img_row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_p = cap_row[j].paragraphs[0]
                        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_run = cap_p.add_run(textwrap.shorten(caption, width=70, placeholder="…"))
                        cap_run.font.size = Pt(8)
                        cap_run.italic = True
                        cap_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                    except Exception:
                        continue
            doc.add_page_break()

    # ── Web Sources ──────────────────────────────────────────────────────────
    if web_on and web_intel.get("results"):
        _add_horizontal_rule(doc)
        wh = doc.add_heading("Web Sources & Further Reading", level=1)
        _set_docx_heading_color(wh, "22D3EE")
        for i, r in enumerate(web_intel["results"], 1):
            p = doc.add_paragraph()
            run_title = p.add_run(f"{i}. {r['title']} — ")
            run_title.bold = True
            run_url = p.add_run(r["url"])
            run_url.font.color.rgb = RGBColor(0x0E, 0x74, 0x90)
        doc.add_page_break()

    # ── Sources ───────────────────────────────────────────────────────────────
    _add_horizontal_rule(doc)
    sh = doc.add_heading("Source Documents", level=1)
    _set_docx_heading_color(sh, "6C63FF")
    for i, name in enumerate(doc_names, 1):
        doc.add_paragraph(f"{i}. {name}", style="List Number")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# PUBLIC ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

def generate_full_report(
    username: str,
    session_id: str,
    settings: dict,
    selected_docs: list[str],
    fmt: str = "pdf",
    use_web: bool = False,
) -> tuple[bytes, str, str, Optional[str]]:
    """
    Generate a full report for the given user's documents.

    Parameters
    ----------
    username      : logged-in user
    session_id    : current chat session (for LLM settings)
    settings      : model/temperature/system_prompt dict from get_session_settings()
    selected_docs : list of filenames to include (empty = all)
    fmt           : "pdf" or "docx"
    use_web       : if True, enrich the report with ONE Tavily web search call
                    (synthesized answer + sources + images) — this does NOT
                    add any extra LLM calls; the same 4 section calls are reused.

    Returns
    -------
    (file_bytes, filename, mime, web_warning)

    web_warning is None when use_web=False, or when web research succeeded.
    It is a short human-readable string when use_web=True but the Tavily call
    failed or returned nothing (e.g. missing TAVILY_API_KEY) — the report is
    still generated (doc-only) and this string explains why. The report file
    itself also shows this same warning on its cover page, so nothing is lost
    if the caller ignores this value.

    NOTE — this is a 4-tuple now (previously 3). Update the Flask route:
        file_bytes, filename, mime, web_warning = generate_full_report(...)
        resp = send_file(io.BytesIO(file_bytes), mimetype=mime,
                          download_name=filename, as_attachment=True)
        if web_warning:
            resp.headers["X-Report-Web-Warning"] = web_warning
        return resp
    """
    # 1. Pull document text from ChromaDB
    doc_map = fetch_document_chunks(username, selected_docs)
    if not doc_map:
        raise ValueError("No documents found. Please upload and index files first.")

    # 1B. Optional single web-research call (Tavily) — cheap, no LLM cost
    web_intel = fetch_web_intel(doc_map) if use_web else None
    web_warning = web_intel.get("error") if (web_intel and not web_intel.get("enabled")) else None

    # 2. Generate all sections via LLM (web content folded into same 4 calls)
    sections = generate_all_sections(doc_map, session_id, settings, web_intel)

    # 3. Render to chosen format
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    if fmt == "docx":
        file_bytes = build_docx_report(username, doc_map, sections, selected_docs, web_intel)
        filename   = f"nexora_report_{username}_{timestamp}.docx"
        mime       = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        file_bytes = build_pdf_report(username, doc_map, sections, selected_docs, web_intel)
        filename   = f"nexora_report_{username}_{timestamp}.pdf"
        mime       = "application/pdf"

    return file_bytes, filename, mime, web_warning
